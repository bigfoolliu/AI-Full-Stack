import tempfile
from pathlib import Path

from docx import Document as DocxDocument
from fpdf import FPDF

from app.services.document_parser import parse_document


def create_test_pdf(path: str):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=14)
    pdf.cell(text="Hello from PDF - nihao shijie", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(text="dierhang neirong", new_x="LMARGIN", new_y="NEXT")
    pdf.output(path)


def create_test_docx(path: str):
    doc = DocxDocument()
    doc.add_paragraph("Hello from Word - nihao shijie")
    doc.add_paragraph("dierhang neirong")
    doc.save(path)


def create_test_txt(path: str):
    Path(path).write_text("Hello from TXT - nihao shijie\ndierhang neirong", encoding="utf-8")


def test_parser():
    with tempfile.TemporaryDirectory() as tmpdir:
        pdf_path = Path(tmpdir) / "test.pdf"
        docx_path = Path(tmpdir) / "test.docx"
        txt_path = Path(tmpdir) / "test.txt"

        create_test_pdf(str(pdf_path))
        create_test_docx(str(docx_path))
        create_test_txt(str(txt_path))

        result_pdf = parse_document(str(pdf_path), "pdf")
        result_docx = parse_document(str(docx_path), "docx")
        result_txt = parse_document(str(txt_path), "txt")

        assert "Hello from PDF" in result_pdf, f"PDF 解析失败: {result_pdf}"
        assert "nihao shijie" in result_pdf, f"PDF 中文解析失败: {result_pdf}"

        assert "Hello from Word" in result_docx, f"Word 解析失败: {result_docx}"
        assert "nihao shijie" in result_docx, f"Word 中文解析失败: {result_docx}"

        assert "Hello from TXT" in result_txt, f"TXT 解析失败: {result_txt}"
        assert "nihao shijie" in result_txt, f"TXT 中文解析失败: {result_txt}"

        print("=" * 40)
        print("PDF 解析结果:")
        print(result_pdf)
        print("=" * 40)
        print("Word 解析结果:")
        print(result_docx)
        print("=" * 40)
        print("TXT 解析结果:")
        print(result_txt)
        print("=" * 40)

        try:
            parse_document(str(pdf_path), "unsupported")
            assert False, "应抛出异常但未抛出"
        except ValueError as e:
            assert "不支持的文件类型" in str(e)

        print("所有测试通过！")
        print("不支持类型异常测试通过！")


if __name__ == "__main__":
    test_parser()
